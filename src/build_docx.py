#!/usr/bin/env python3
"""Build baseline CV as a .docx Word document with two-column layout."""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# --- Color palette (matching current CV accent) ---
ACCENT = RGBColor(0x3B, 0xC0, 0xE3)  # teal/cyan
DARK_BG = RGBColor(0x37, 0x47, 0x4F)  # charcoal header
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_DARK = RGBColor(0x26, 0x32, 0x38)
TEXT_MID = RGBColor(0x45, 0x5A, 0x64)
TEXT_LIGHT = RGBColor(0x78, 0x90, 0x9C)
SIDEBAR_BG = "ECEFF1"
DOT_FILLED = "\u25CF"  # ●
DOT_EMPTY = "\u25CB"   # ○

OUTPUT_FILE = Path(__file__).parent.parent / "Luis_Vela_CV.docx"


# ─── Utility helpers ───────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set individual cell margins in EMUs."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_existing = tcPr.find(qn('w:tcMar'))
    if tcMar_existing is not None:
        tcPr.remove(tcMar_existing)
    tcPr.append(tcMar)


def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:start w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:end w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_paragraph_spacing(para, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def add_run(para, text, size=8, color=TEXT_MID, bold=False, italic=False, font_name="Calibri"):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    return run


def add_sidebar_heading(cell, text):
    """Add a section heading in the sidebar."""
    para = cell.add_paragraph()
    set_paragraph_spacing(para, before=6, after=2)
    add_run(para, text, size=8.5, color=ACCENT, bold=True)
    # Add accent underline
    para2 = cell.add_paragraph()
    set_paragraph_spacing(para2, before=0, after=3)
    run = para2.add_run()
    run.font.size = Pt(1)
    # Use a bottom border on the paragraph instead
    pPr = para.paragraph_format.element
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="3BC0E3"/>'
        f'</w:pBdr>'
    )
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)
    # Remove the empty spacer paragraph
    cell._tc.remove(para2._element)


def add_main_heading(cell, text, icon=""):
    """Add a section heading in the main column."""
    para = cell.add_paragraph()
    set_paragraph_spacing(para, before=6, after=3)
    if icon:
        add_run(para, icon + " ", size=8, color=ACCENT)
    add_run(para, text, size=9, color=ACCENT, bold=True)
    # Bottom border
    pPr = para.paragraph_format.element
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="3BC0E3"/>'
        f'</w:pBdr>'
    )
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def add_skill_dots(cell, name, filled, total=5, sub_text=None):
    """Add a skill row with dot rating."""
    para = cell.add_paragraph()
    set_paragraph_spacing(para, before=0, after=1, line=11)
    add_run(para, name, size=7, color=TEXT_MID)
    dots = "  " + (DOT_FILLED + " ") * filled + (DOT_EMPTY + " ") * (total - filled)
    run_dots = add_run(para, dots.rstrip(), size=6.5, color=ACCENT)
    # Color empty dots differently — we'll just use one color for simplicity
    # The filled/empty distinction is visual from the glyph itself
    if sub_text:
        para2 = cell.add_paragraph()
        set_paragraph_spacing(para2, before=0, after=1, line=10)
        add_run(para2, sub_text, size=6, color=TEXT_LIGHT, italic=True)


def add_bullet_item(cell, text, size=7, icon="\u25B8"):
    """Add a bullet point item."""
    para = cell.add_paragraph()
    set_paragraph_spacing(para, before=0, after=2, line=11)
    add_run(para, icon + " ", size=6.5, color=ACCENT)
    add_run(para, text, size=size, color=TEXT_MID)


def add_experience_entry(cell, title, date_range, location, bullets):
    """Add a professional experience or education entry."""
    # Title line
    para_title = cell.add_paragraph()
    set_paragraph_spacing(para_title, before=2, after=1, line=12)
    add_run(para_title, title, size=8, color=TEXT_DARK, bold=True)

    # Date and location
    para_meta = cell.add_paragraph()
    set_paragraph_spacing(para_meta, before=0, after=2, line=10)
    add_run(para_meta, date_range, size=6.5, color=TEXT_LIGHT)
    add_run(para_meta, "  |  ", size=6.5, color=TEXT_LIGHT)
    add_run(para_meta, location, size=6.5, color=TEXT_LIGHT)

    # Bullet points
    for bullet in bullets:
        add_bullet_item(cell, bullet, size=6.8)


# ─── Content builders ──────────────────────────────────────────────────────────

def build_header(doc):
    """Build the dark charcoal header with name and headline."""
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(header_table)

    # Set full width
    tbl = header_table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    existing_w = tblPr.find(qn('w:tblW'))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)

    cell = header_table.rows[0].cells[0]
    set_cell_shading(cell, "37474F")
    set_cell_margins(cell, top=180, bottom=140, left=240, right=240)

    # Name
    para_name = cell.paragraphs[0]
    set_paragraph_spacing(para_name, before=0, after=2)
    add_run(para_name, "LUIS VELA VELA PhD", size=18, color=WHITE, bold=True)

    # Headline
    para_headline = cell.add_paragraph()
    set_paragraph_spacing(para_headline, before=0, after=0)
    add_run(
        para_headline,
        "Sr. AI Scientist & ML Team Lead  |  Computational Physicist",
        size=9, color=RGBColor(0xCF, 0xD8, 0xDC), bold=False
    )


def build_sidebar(cell):
    """Build the left sidebar content."""
    set_cell_shading(cell, SIDEBAR_BG)
    set_cell_margins(cell, top=140, bottom=100, left=170, right=120)

    # Remove default empty paragraph
    if cell.paragraphs:
        cell.paragraphs[0].clear()

    # ── CONTACT ──
    add_sidebar_heading(cell, "CONTACT")

    contact_items = [
        ("\u2709", "vela.vela.luis@gmail.com"),
        ("\u260E", "+352 661 678965"),
        ("\u25CF", "Luxembourg, Luxembourg"),
        ("\u2302", "OnlineResume"),
        ("GH", "@luiservela"),
        ("in", "Luis Vela Vela"),
    ]
    for icon, value in contact_items:
        para = cell.add_paragraph()
        set_paragraph_spacing(para, before=0, after=1.5, line=10)
        add_run(para, icon + "  ", size=7, color=ACCENT)
        add_run(para, value, size=6.8, color=TEXT_MID)

    # ── CORE COMPETENCIES ──
    add_sidebar_heading(cell, "CORE COMPETENCIES")

    # AI/ML & Deep Learning
    cat_para = cell.add_paragraph()
    set_paragraph_spacing(cat_para, before=2, after=2, line=10)
    add_run(cat_para, "AI/ML & Deep Learning", size=7, color=ACCENT, italic=True)

    add_skill_dots(cell, "ML/DL Frameworks", 5, sub_text="(TensorFlow, PyTorch)")
    add_skill_dots(cell, "AI Weather Models", 5, sub_text="(Aurora, AIFS, FourCastNet, GraphCast)")
    add_skill_dots(cell, "AI-Assisted Dev", 4, sub_text="(Claude Code, MCP integrations)")
    add_skill_dots(cell, "Experiment Tracking", 4, sub_text="(MLflow)")
    add_skill_dots(cell, "Probabilistic Forecasting", 4, sub_text="(CRPS, ensemble methods)")

    # NWP & Atmospheric Science
    cat_para = cell.add_paragraph()
    set_paragraph_spacing(cat_para, before=3, after=2, line=10)
    add_run(cat_para, "NWP & Atmospheric Science", size=7, color=ACCENT, italic=True)

    add_skill_dots(cell, "NWP Pipelines", 5)
    add_skill_dots(cell, "GRIB2 Expertise", 5)
    add_skill_dots(cell, "ERA5 Reanalysis", 5)
    add_skill_dots(cell, "S2S Forecasting", 4, sub_text="(EC46, weather regimes)")
    add_skill_dots(cell, "Renewable Energy Fcst", 4)

    # Programming & Development
    cat_para = cell.add_paragraph()
    set_paragraph_spacing(cat_para, before=3, after=2, line=10)
    add_run(cat_para, "Programming & Development", size=7, color=ACCENT, italic=True)

    add_skill_dots(cell, "Python", 5, sub_text="(xarray, NumPy, Pandas, Scikit-learn, Matplotlib, Cartopy)")
    add_skill_dots(cell, "GPU/Distributed", 4, sub_text="(CuPy, CuDF, CuML, Dask, RAPIDS, Numba)")
    add_skill_dots(cell, "Bash", 4)
    add_skill_dots(cell, "FORTRAN", 3)

    # Infrastructure & DevOps
    cat_para = cell.add_paragraph()
    set_paragraph_spacing(cat_para, before=3, after=2, line=10)
    add_run(cat_para, "Infrastructure & DevOps", size=7, color=ACCENT, italic=True)

    add_skill_dots(cell, "HPC Operations", 5, sub_text="(Lmod, SSH, GPU clusters \u2014 A100, H200)")
    add_skill_dots(cell, "Version Control", 5, sub_text="(Git, Bitbucket, GitHub, GitLab)")
    add_skill_dots(cell, "CI/CD & Quality", 4, sub_text="(pylint, mypy, pytest, coverage)")
    add_skill_dots(cell, "Env Management", 4, sub_text="(Conda)")
    add_skill_dots(cell, "Cloud (AWS)", 3)

    # ── LANGUAGES ──
    add_sidebar_heading(cell, "LANGUAGES")

    add_skill_dots(cell, "Spanish", 5, sub_text="Native")
    add_skill_dots(cell, "English", 5, sub_text="Fluent / Professional")
    add_skill_dots(cell, "French", 4, sub_text="Professional")
    add_skill_dots(cell, "Serbian", 3, sub_text="Conversational")
    add_skill_dots(cell, "Czech", 3, sub_text="Conversational")
    add_skill_dots(cell, "German", 2, sub_text="Working knowledge")

    # ── KEY ATTRIBUTES ──
    add_sidebar_heading(cell, "KEY ATTRIBUTES")

    attributes = [
        "Scientific rigor & reproducibility",
        "Team builder & engineering culture advocate",
        "Production ML delivery focus",
        "Cross-functional communicator",
    ]
    for attr in attributes:
        para = cell.add_paragraph()
        set_paragraph_spacing(para, before=0, after=1.5, line=10)
        add_run(para, "\u25B8 ", size=6.5, color=ACCENT)
        add_run(para, attr, size=6.8, color=TEXT_MID, bold=True)


def build_main(cell):
    """Build the right main column content."""
    set_cell_margins(cell, top=140, bottom=100, left=170, right=200)

    # Remove default empty paragraph
    if cell.paragraphs:
        cell.paragraphs[0].clear()

    # ── PROFESSIONAL SUMMARY ──
    add_main_heading(cell, "PROFESSIONAL SUMMARY")
    para = cell.add_paragraph()
    set_paragraph_spacing(para, before=1, after=4, line=11)
    add_run(
        para,
        "Sr. AI Scientist with a PhD in Computational Physics, currently leading an ML "
        "engineering team building AI-powered weather forecasting systems at Spire Global. "
        "Combines deep scientific expertise in numerical weather prediction and atmospheric "
        "science with production ML engineering and team leadership. Proven track record "
        "across research (Amazon), HPC solutions engineering (LuxProvide), and operational "
        "AI deployment (Spire).",
        size=7, color=TEXT_MID
    )

    # ── PROFESSIONAL EXPERIENCE ──
    add_main_heading(cell, "PROFESSIONAL EXPERIENCE", icon="\u2699")

    # --- Spire ---
    add_experience_entry(
        cell,
        title="Sr. AI Weather Scientist \u2014 Spire Global, Luxembourg",
        date_range="Nov 2022 \u2013 Present",
        location="Luxembourg",
        bullets=[
            "Manage a team of 2 ML engineers. Established team-wide engineering practices "
            "including experiment tracking (MLflow), shared repositories, code review "
            "standards, and documentation workflows.",
            "Architected and deployed production infrastructure for multiple AI-based NWP "
            "models (AIFS, Aurora). Built model comparison and verification tooling for "
            "forecast intercomparison across ensemble statistics, weather regimes, and "
            "spatial fields.",
            "Developed subseasonal-to-seasonal (S2S) forecasting pipelines including "
            "weather regime classification and ensemble probability tracking.",
            "Deep expertise in GRIB2 format handling, ERA5 reanalysis data processing "
            "(accumulated vs. instantaneous variables, temporal aggregation), and "
            "multi-model data pipelines.",
            "Pioneered adoption of Claude Code for AI-assisted software development on HPC "
            "infrastructure, achieving 6\u201312x speedup in model repository scaffolding.",
            "Contributed to weather-to-energy production forecast pipelines for renewable "
            "energy applications.",
            "Presented technical solutions to internal and external stakeholders ensuring "
            "successful project delivery.",
        ],
    )

    # --- LuxProvide ---
    add_experience_entry(
        cell,
        title="Sr. Solutions Engineer \u2014 LuxProvide, Luxembourg",
        date_range="Feb 2021 \u2013 Nov 2022",
        location="Luxembourg",
        bullets=[
            "Conducted technical discovery and designed custom HPC/AI solutions for "
            "clients during sales cycles.",
            "Managed client relationships from initial assessment through deployment.",
            "Led pre-sales technical presentations leading to HPC solution contracts.",
            "Evaluated GPU computing platforms including H200 infrastructure.",
        ],
    )

    # --- Amazon ---
    add_experience_entry(
        cell,
        title="Research Scientist \u2014 Amazon, Luxembourg",
        date_range="Oct 2019 \u2013 Jan 2021",
        location="Luxembourg",
        bullets=[
            "Applied advanced ML and statistical methods to deliver actionable business "
            "insights with direct impact.",
        ],
    )

    # ── EDUCATION ──
    add_main_heading(cell, "EDUCATION", icon="\U0001F393")

    add_experience_entry(
        cell,
        title="PhD in Computational Physics",
        date_range="Sep 2013 \u2013 Feb 2019",
        location="UC3M, Madrid | UGent, Ghent",
        bullets=[
            "Specialized in computational methods for complex physical systems. "
            "Developed algorithms for HPC environments.",
        ],
    )

    add_experience_entry(
        cell,
        title="MSc in Plasma Physics",
        date_range="Sep 2011 \u2013 Jul 2013",
        location="UC3M, Madrid | UGent, Ghent",
        bullets=[
            "Statistical analysis and modeling of complex dynamic systems.",
        ],
    )

    add_experience_entry(
        cell,
        title="BSc in Physics",
        date_range="Sep 2007 \u2013 Jul 2010",
        location="Charles University, Prague",
        bullets=[
            "Foundation in computational physics and simulation methods.",
        ],
    )

    # ── ACHIEVEMENTS & RECOGNITION ──
    add_main_heading(cell, "ACHIEVEMENTS & RECOGNITION", icon="\U0001F3C6")

    achievements = [
        "Outstanding Colombian Abroad \u2014 Award by the Colombian Government",
        "Summa Cum Laude \u2014 PhD Thesis",
        "Greatest Distinction \u2014 2013 Erasmus Mundus Master",
        "UNESCO Fellowship \u2014 Bachelor Studies Scholarship",
    ]
    for ach in achievements:
        add_bullet_item(cell, ach, size=6.8, icon="\U0001F3C6")

    # ── SELECTED PUBLICATIONS ──
    add_main_heading(cell, "SELECTED PUBLICATIONS", icon="\U0001F4DA")

    publications = [
        "Magneto-hydrodynamical nonlinear simulations of magnetically confined plasmas "
        "using smooth particle hydrodynamics (SPH)",
        "A positioning algorithm for SPH in smoothly curved geometries",
        "ALARIC: An algorithm for constructing arbitrarily complex initial density "
        "distributions with low particle noise for SPH/SPMHD applications",
    ]
    for pub in publications:
        add_bullet_item(cell, pub, size=6.5, icon="\U0001F4C4")


# ─── Main builder ──────────────────────────────────────────────────────────────

def build_cv():
    """Build the complete CV document."""
    doc = Document()

    # --- Page setup: A4, narrow margins ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # --- Default font ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(8)
    font.color.rgb = TEXT_DARK
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(11)

    # ── HEADER ──
    build_header(doc)

    # ── BODY: two-column table ──
    body_table = doc.add_table(rows=1, cols=2)
    body_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(body_table)

    # Set full width
    tbl = body_table._tbl
    tblPr = tbl.tblPr
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    existing_w = tblPr.find(qn('w:tblW'))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblPr.append(tblW)

    # Column widths: ~34% sidebar, ~66% main
    sidebar_width = Cm(7.1)   # ~34% of 21cm
    main_width = Cm(13.9)     # ~66% of 21cm

    sidebar_cell = body_table.rows[0].cells[0]
    main_cell = body_table.rows[0].cells[1]

    # Set column widths
    sidebar_cell.width = sidebar_width
    main_cell.width = main_width

    # Also set grid column widths for better compatibility
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn('w:gridCol'))
        if len(gridCols) >= 2:
            gridCols[0].set(qn('w:w'), str(int(sidebar_width.emu / 635)))
            gridCols[1].set(qn('w:w'), str(int(main_width.emu / 635)))

    # Build content
    build_sidebar(sidebar_cell)
    build_main(main_cell)

    # Remove borders from cells
    for row in body_table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)

    return doc


def main():
    output = OUTPUT_FILE if len(sys.argv) < 2 else Path(sys.argv[1])
    print(f"Building {output} ...")
    doc = build_cv()
    doc.save(str(output))
    print(f"Done. Output: {output} ({output.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
